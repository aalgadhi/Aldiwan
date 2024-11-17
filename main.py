import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
from os import system
import tkinter as tk
from tkinter import messagebox
from tkinter import PhotoImage

harakat = [character for character in "ًٌٍَُِّْ"]  # All of the harakat

def add_mosahmat_to_end(doc, mosahmat_items):
    # Add the mosahmat content at the end of the document
    doc.add_paragraph("\n:المعاني", style="Heading 1").alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    for item in mosahmat_items:
        h2_text = item.find("h2").get_text().strip() if item.find("h2") else ""
        h4_text = item.find("h4").get_text().strip() if item.find("h4") else ""
        h2_text = h2_text.replace(".", "").replace("\n", "") #.replace("\r", "")
        h4_text = h4_text.replace(".", "").replace("\n", "") #.replace("\r", "")

        if h2_text:
            paragraph_h2 = doc.add_paragraph(h2_text, style="Heading 2")
            paragraph_h2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if h4_text:
            paragraph_h2 = doc.add_paragraph(h4_text)
            paragraph_h2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            paragraph_h2.style.font.size = Pt(12)  # Adjust font size for h2

def baitPrinter(alsader, alajs):
    return alsader.strip(), alajs.strip()


def generate_qasida_pdf():
    # Get the URL from the entry field
    url = url_entry.get()

    # Check if URL is empty
    if not url:
        messagebox.showwarning("Input Error", "Please enter a URL.")
        return

    # Send a GET request to fetch the HTML content of the page
    response = requests.get(url)

    # Check if the request was successful
    if response.status_code == 200:
        # Parse the HTML content
        soup = BeautifulSoup(response.text, "html.parser")

        # Create a new Document
        doc = Document()

        # Qasida information
        qasida_info = soup.find("h2").get_text().split("»")
        matlaa_alqasida = qasida_info[-1].strip()  # First line of the poetry
        alshaer = qasida_info[-2].strip()  # The poet
        matlaa_alqasida_splitted = "_".join(map(str, matlaa_alqasida.split()))
        alshaer_splitted = "_".join(map(str, alshaer.split()))

        # Add poet name as heading
        doc.add_heading(alshaer, level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        bahr_alqasida = soup.find_all("div", class_="col-6 col-md-3")[2].get_text().strip()
        paragraph_albahr = doc.add_paragraph(":" + "من " + bahr_alqasida, style="Heading 2")
        paragraph_albahr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        # Locate the div with id "poem_content" and find all h3 tags within it
        poem_content = soup.find("div", id="poem_content")
        if poem_content:
            h3_tags = poem_content.find_all("h3")

            # Create a table with two columns for "alsader" and "alajs"
            table = doc.add_table(rows=0, cols=2)

            # Loop through h3 tags in pairs to form verses
            for i in range(int(len(h3_tags) / 2)):
                alsader = h3_tags[2 * i].get_text().strip()  # First half of the line
                alajs = h3_tags[2 * i + 1].get_text().strip()  # Second half of the line

                # Add a row to the table
                row_cells = table.add_row().cells
                row_cells[1].text = alsader  # Print "alsader" in the first column
                row_cells[0].text = alajs  # Print "alajs" in the second column

                # Set alignment and font size for both cells
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        paragraph.style.font.size = Pt(14)
            
            if include_mosahamat.get():
                mosahmat_items = soup.find_all("div", class_="mosahmat_item")
                if mosahmat_items:
                    add_mosahmat_to_end(doc, mosahmat_items)

            # Generate file names
            system(f"mkdir qasidas")
            file_name = f"{matlaa_alqasida_splitted}_{alshaer_splitted}"
            docx_filename = f"qasidas/{file_name}.docx"
            pdf_filename = f"qasidas/{file_name}.pdf"

            # Save the document as a .docx file
            doc.save(docx_filename)

            # Convert the .docx file to a PDF
            convert(docx_filename, pdf_filename)

            # Open the PDF file
            system(f"start {pdf_filename}")
            messagebox.showinfo("Success", f"'{pdf_filename}' حفظت القصيدة بالملف: ")
        else:
            messagebox.showerror("Content Error", "القصيدة يجب أن تكون من موقع الديوان")
    else:
        messagebox.showerror("Request Error", f"Failed to retrieve the page. Status code: {response.status_code}")


if __name__ == "__main__":
    # Create the GUI window
    root = tk.Tk()
    root.title("Qasida PDF Generator")
    include_mosahamat = tk.BooleanVar(value=True)
    # Set the window icon to "logo.ico" for both title bar and taskbar
    root.iconbitmap("logo.ico")

    # URL input label and entry field
    url_entry = tk.Entry(root, width=50)
    url_entry.grid(row=0, column=0, padx=10, pady=10)

    tk.Label(root, text=":الصق الرابط من موقع الديوان").grid(row=0, column=1, padx=10, pady=10)
    include_mosahamat_checkbox = tk.Checkbutton(
        root, text="إضافة معاني القصيدة", variable=include_mosahamat
    )
    include_mosahamat_checkbox.grid(row=1, column=0, columnspan=2, pady=5)

    # Generate PDF button
    generate_button = tk.Button(root, text="ابدأ", command=generate_qasida_pdf)
    generate_button.grid(row=2, column=0, columnspan=2, pady=10)

    # Run the GUI
    root.mainloop()